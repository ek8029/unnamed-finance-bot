"""
Render docs/mucker-reference.md into a styled .docx on the Desktop.

Living workflow: edit the Markdown master (docs/mucker-reference.md) over time,
then re-run this to refresh the Word doc.

    python scripts/render-mucker-doc.py

Output: C:\\Users\\Evan\\Desktop\\Helm Mucker Reference.docx
"""
import os
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

GOLD = RGBColor(0xB8, 0x86, 0x0B)      # readable gold on white (brand E6B94D is too light)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "..", "docs", "mucker-reference.md")
DEFAULT_OUT = os.path.join(os.path.expanduser("~"), "Desktop", "Helm Mucker Reference.docx")


def add_runs(paragraph, text):
    """Render inline **bold** segments as bold runs; everything else plain."""
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "":
            continue
        run = paragraph.add_run(seg)
        run.font.color.rgb = DARK
        if i % 2 == 1:  # odd segments were inside ** **
            run.bold = True


def main():
    out = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_OUT
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK

    for raw in lines:
        line = raw.rstrip()

        if not line.strip():
            continue

        if line.strip() == "---":
            # thin separator: an empty spacer paragraph
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:].strip())
            r.bold = True
            r.font.size = Pt(22)
            r.font.color.rgb = GOLD
            continue

        if line.startswith("## "):
            doc.add_paragraph()
            p = doc.add_paragraph()
            r = p.add_run(line[3:].strip())
            r.bold = True
            r.font.size = Pt(15)
            r.font.color.rgb = GOLD
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:].strip())
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = DARK
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(line[2:].strip())
            r.italic = True
            r.font.color.rgb = GREY
            continue

        # bullets: "- " top level, "  - " nested
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent, content = m.group(1), m.group(2)
            style = "List Bullet 2" if len(indent) >= 2 else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            add_runs(p, content)
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_runs(p, line.strip())

    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
